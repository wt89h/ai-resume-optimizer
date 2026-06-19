from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

class ExportService:
    async def export(self, modification_id: str, format: str) -> tuple[bytes, str]:
        """导出简历"""
        if format == "pdf":
            return await self.export_pdf(modification_id)
        else:
            return await self.export_docx(modification_id)
    
    async def export_pdf(self, modification_id: str) -> tuple[bytes, str]:
        """导出PDF格式"""
        # 这里可以使用reportlab生成PDF
        # 目前返回模拟PDF
        filename = f"resume_{modification_id}.pdf"
        content = b"%PDF-1.4 mock pdf content"
        return content, filename
    
    async def export_docx(self, modification_id: str) -> tuple[bytes, str]:
        """导出Word格式"""
        doc = Document()
        
        # 设置标题
        title = doc.add_heading('个人简历', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加个人信息
        doc.add_heading('个人信息', level=1)
        doc.add_paragraph('姓名：张三')
        doc.add_paragraph('邮箱：zhangsan@example.com')
        doc.add_paragraph('电话：13800138000')
        
        # 添加工作经历
        doc.add_heading('工作经历', level=1)
        doc.add_paragraph('ABC科技有限公司 | 数据分析师 | 2022.07 - 至今')
        doc.add_paragraph('负责数据分析项目，使用Python进行数据处理，通过数据挖掘优化业务流程，提升效率20%')
        
        # 添加技能
        doc.add_heading('专业技能', level=1)
        doc.add_paragraph('• Python数据分析（Pandas、NumPy）')
        doc.add_paragraph('• SQL数据库查询')
        doc.add_paragraph('• 数据可视化（Tableau、Power BI）')
        
        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"resume_{modification_id}.docx"
        return buffer.read(), filename
